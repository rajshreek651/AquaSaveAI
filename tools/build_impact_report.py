"""Create the AquaSave AI impact-assessment report as a formatted DOCX."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path(__file__).resolve().parents[1]
REPORT = json.loads((PROJECT / "data" / "analysis_report.json").read_text(encoding="utf-8"))
OUTPUT = PROJECT / "outputs" / "AquaSaveAI_Impact_Assessment_Report.docx"

NAVY = "063B49"
TEAL = "0D9488"
MINT = "D9F5ED"
PALE = "F2F7F6"
MUTED = "60767B"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.margin_top = 80
            cell.margin_bottom = 80


def style_run(run, size: float = 11, color: str = NAVY, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(doc, text: str, size: float = 11, color: str = NAVY, bold: bool = False, after: float = 6) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    style_run(paragraph.add_run(text), size, color, bold)


def add_heading(doc, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    paragraph.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    style_run(paragraph.add_run(text), 16 if level == 1 else 13, TEAL, True)


def add_metric_strip(doc, metrics: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    set_table_geometry(table, [9360 // len(metrics)] * len(metrics))
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, MINT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        style_run(p.add_run(value), 17, NAVY, True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        style_run(p2.add_run(label), 8.5, MUTED)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    style_run(p.add_run(text), 11, NAVY)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        style_run(p.add_run(header), 9.5, WHITE, True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            set_cell_shading(cell, PALE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            style_run(p.add_run(value), 9.5, NAVY)


def build() -> None:
    OUTPUT.parent.mkdir(exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run("AquaSave AI | Impact Assessment"), 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(footer.add_run("Simulation-based prototype | SDG 6: Clean Water and Sanitation"), 8.5, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run("IMPACT ASSESSMENT REPORT"), 10, TEAL, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run("AquaSave AI"), 27, NAVY, True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    style_run(p.add_run("A lightweight water-wastage detection and conservation simulator"), 13, MUTED)

    projected_monthly = REPORT["potential_waste_litres"] / 14 * 30
    projected_cost = projected_monthly * 0.03
    projected_energy = projected_monthly / 1000 * 0.5
    projected_co2 = projected_energy * 0.7
    add_metric_strip(doc, [
        (f"{REPORT['potential_waste_litres']:.1f} L", "potential waste in 14 days"),
        (f"{projected_monthly:.1f} L", "projected monthly water saving"),
        (f"Rs {projected_cost:.2f}", "projected monthly cost saving"),
        (f"{projected_co2:.3f} kg", "projected monthly CO2 avoided"),
    ])

    add_heading(doc, "1. Executive summary", 1)
    add_text(doc, "AquaSave AI analysed 336 hourly water-use records for a simulated four-person household. It identified two actionable patterns: sustained overnight flow that may indicate a leak and a single period of unusually high garden use. The prototype converts these alerts into understandable actions and estimated sustainability impact.")

    add_heading(doc, "2. What the prototype detected", 1)
    alerts = REPORT["alerts"]
    rows = [[a["type"], a["severity"], a["period"], f"{a['potential_waste_litres']:.2f} L"] for a in alerts]
    add_table(doc, ["Detected pattern", "Priority", "When", "Potential waste"], rows, [2850, 1050, 3260, 2200])
    add_text(doc, "The alert logic uses transparent rules: persistent overnight readings above 8 L/hour indicate a possible continuous leak; a reading at least 10 litres above its hour-specific statistical range indicates unusual use.", 9.5, MUTED, after=8)

    add_heading(doc, "3. Estimated impact", 1)
    add_text(doc, "The values below are scenario estimates. They assume that the detected water wastage is prevented consistently over a 30-day month.")
    impact_rows = [
        ["Water", f"{projected_monthly:.1f} litres/month", "Potential water conserved after corrective action"],
        ["Cost", f"Rs {projected_cost:.2f}/month", "Using the editable prototype tariff of Rs 0.03 per litre"],
        ["Pumping energy", f"{projected_energy:.3f} kWh/month", "Using 0.5 kWh per 1,000 litres pumped"],
        ["Carbon emissions", f"{projected_co2:.3f} kg CO2/month", "Using 0.7 kg CO2 per kWh of electricity"],
    ]
    add_table(doc, ["Impact area", "Projected result", "Interpretation"], impact_rows, [1850, 2600, 4560])

    add_heading(doc, "4. Recommended actions", 1)
    add_bullet(doc, "Inspect toilet flush tanks, taps, and concealed pipe connections after a sustained overnight-flow alert.")
    add_bullet(doc, "Review garden watering duration and avoid watering when demand is already unusually high.")
    add_bullet(doc, "Use the dashboard's conservation simulator to compare leak repair and shower-use reduction scenarios.")

    add_heading(doc, "5. Limitations and scale-up plan", 1)
    add_text(doc, "This is a simulation-based prototype, not a live smart-meter deployment. Its water tariff, pumping-energy factor, and electricity-emission factor are transparent project assumptions and should be replaced with local utility values for a formal assessment. At scale, hourly data can be collected from a smart meter or IoT flow sensor; a larger labelled or Kaggle dataset can then be used to compare the existing rules against an Isolation Forest anomaly-detection model.")

    add_heading(doc, "6. Assignment-roadmap alignment", 1)
    add_table(doc, ["Roadmap stage", "AquaSave AI evidence"], [
        ["Identify challenge", "Water wastage through leaks and unexpected usage"],
        ["Explore and visualise data", "336 hourly records, daily trend, source summaries, and alerts"],
        ["Design smart solution", "Explainable leak and unusual-use detection with conservation scenarios"],
        ["Assess impact", "Water, cost, pumping energy, and CO2 projections"],
        ["Plan to scale", "Smart-meter integration, more real data, optional machine-learning comparison"],
    ], [2700, 6660])

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
